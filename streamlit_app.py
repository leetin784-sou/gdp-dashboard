from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Basic styling
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
style.font.size = Pt(13)

def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)

def add_h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True

def add_para(text):
    doc.add_paragraph(text)

def add_equation_block(text):
    # Word equation objects are complex; we place as monospace-style text block.
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    run.font.size = Pt(11)

add_title("BÁO CÁO KHKTQG (BẢN HOÀN CHỈNH)\nMô hình hỗ trợ quyết định phân luồng cấp cứu dựa trên AI\nvới định lượng rủi ro và độ không chắc chắn")
add_para("Họ và tên: .............................................")
add_para("Trường/Lớp: ............................................")
add_para("Lĩnh vực: Tin học/Y sinh/Trí tuệ nhân tạo ứng dụng")
doc.add_paragraph()

add_h1("TÓM TẮT (ABSTRACT)")
add_para(
"Trong khoa Cấp cứu, quyết định phân luồng (triage) phải được đưa ra trong thời gian rất ngắn và chịu ảnh hưởng mạnh bởi quá tải, "
"mệt mỏi và sự khác nhau về kinh nghiệm cá nhân. Sai sót nguy hiểm nhất không phải là chẩn đoán sai bệnh, mà là bỏ sót bệnh nhân "
"nguy kịch hoặc xếp mức ưu tiên thấp trong những phút đầu. Đề tài đề xuất một mô hình hỗ trợ quyết định phân luồng dựa trên AI theo "
"hướng an toàn: (i) kết hợp luật y khoa (red flags) với mô hình học máy, (ii) thay vì phân loại cứng Đỏ/Vàng/Xanh, hệ thống ước lượng "
"xác suất nguy kịch (risk score) liên tục, và (iii) định lượng độ không chắc chắn (uncertainty) để kích hoạt cơ chế Human-in-the-loop "
"đúng thời điểm. Mô hình được đánh giá trên dữ liệu lâm sàng giả lập theo các chỉ số an toàn như Recall nhóm Đỏ và tỷ lệ bỏ sót nguy kịch, "
"đồng thời đo thời gian suy luận/ca. Kết quả kỳ vọng: giảm bỏ sót nguy kịch và chuẩn hóa quyết định dưới áp lực cao trong khi bác sĩ giữ "
"quyền quyết định cuối cùng."
)
doc.add_paragraph()

add_h1("1. GIỚI THIỆU")
add_h2("1.1. Bối cảnh và tính cấp thiết")
add_para(
"Phân luồng cấp cứu (triage) là bước quyết định thứ tự ưu tiên xử trí. Trong thực tế, cùng một tình huống lâm sàng có thể được đánh giá "
"khác nhau giữa các nhân viên y tế, đặc biệt khi khoa cấp cứu quá tải và thời gian đánh giá ban đầu chỉ vài chục giây. Khi đó, lỗi nguy hiểm "
"nhất là bỏ sót ca nguy kịch không điển hình (vital signs lệch nhẹ hoặc triệu chứng mờ) dẫn đến trì hoãn xử trí."
)
add_h2("1.2. Vấn đề nghiên cứu")
add_para(
"Vấn đề đặt ra: Làm thế nào để giảm sai lệch và tăng tính nhất quán của quyết định phân luồng trong 30–60 giây đầu, "
"đồng thời giảm tỷ lệ bỏ sót ca nguy kịch, mà vẫn giữ bác sĩ là người quyết định cuối cùng?"
)
add_h2("1.3. Mục tiêu và phạm vi")
add_para(
"Mục tiêu: Xây dựng mô hình hỗ trợ phân luồng 3 mức (Đỏ/Vàng/Xanh) dựa trên (i) luật y khoa để phát hiện red flags, "
"(ii) mô hình AI ước lượng risk score, và (iii) uncertainty để kích hoạt Human-in-the-loop. Phạm vi: hỗ trợ phân luồng ưu tiên, "
"không chẩn đoán bệnh cuối cùng, không kê đơn/điều trị."
)
doc.add_paragraph()

add_h1("2. CƠ SỞ LÝ THUYẾT")
add_h2("2.1. Phân luồng cấp cứu và red flags")
add_para(
"Triage là phân loại mức ưu tiên dựa trên dấu hiệu sinh tồn và triệu chứng. Red flags là các dấu hiệu nguy kịch cần ưu tiên xử trí ngay "
"(ví dụ: SpO₂ rất thấp, tụt huyết áp, rối loạn ý thức, khó thở nặng...)."
)
add_h2("2.2. Bài toán phân loại và ước lượng rủi ro")
add_para(
"Thay vì chỉ dự đoán nhãn (Đỏ/Vàng/Xanh), đề tài ưu tiên ước lượng xác suất nguy kịch P(critical) để bác sĩ hiểu mức độ nguy hiểm liên tục. "
"Điều này phù hợp với decision science: quyết định dựa trên rủi ro và hậu quả sai."
)
add_h2("2.3. Khái niệm Uncertainty trong mô hình dự đoán")
add_para(
"Uncertainty là mức độ mô hình không chắc chắn với dự đoán. Trong y khoa, mô hình an toàn cần biết khi nào không nên đưa ra khuyến nghị mạnh "
"và khi nào phải yêu cầu bác sĩ đánh giá kỹ."
)
doc.add_paragraph()

add_h1("3. CÔNG TRÌNH LIÊN QUAN")
add_para(
"Các hệ thống hỗ trợ quyết định lâm sàng (CDSS) thường cung cấp kiến thức/phác đồ hoặc hỗ trợ chẩn đoán. Nhiều hệ thống AI y tế tối ưu hóa "
"độ chính xác kỹ thuật (accuracy/AUC hoặc Dice trong ảnh y khoa). Tuy nhiên, các hệ thống này thường không tối ưu cho quyết định tức thời "
"trong cấp cứu và hiếm khi báo uncertainty một cách rõ ràng cho cơ chế Human-in-the-loop."
)
doc.add_paragraph()

add_h1("4. ĐÓNG GÓP VÀ TÍNH ĐỘT PHÁ")
add_para("Đóng góp chính của đề tài:")
doc.add_paragraph("1) Chuyển từ phân luồng cứng sang định lượng rủi ro (risk score 0–100%).")
doc.add_paragraph("2) Tích hợp uncertainty để AI biết khi nào không chắc và kích hoạt Human-in-the-loop đúng thời điểm.")
doc.add_paragraph("3) Đặt trọng tâm đánh giá theo chỉ số an toàn: Recall nhóm Đỏ và tỷ lệ bỏ sót nguy kịch.")
doc.add_paragraph("4) Thiết kế kiến trúc triển khai nhanh cho 30–60 giây đầu, bác sĩ quyết định cuối cùng.")
add_para(
"Tính đột phá không nằm ở việc phát minh thuật toán mới, mà ở cách dùng AI an toàn hơn: định lượng rủi ro + độ không chắc chắn để giảm "
"sai lệch quyết định dưới áp lực cao."
)
doc.add_paragraph()

add_h1("5. PHƯƠNG PHÁP NGHIÊN CỨU")
add_h2("5.1. Đặc trưng đầu vào")
add_para("Mỗi ca gồm các biến tối thiểu (có thể mở rộng): Age, HR, SBP, SpO₂, RR, AVPU/tri giác, nhóm triệu chứng chính.")
add_h2("5.2. Thiết kế dữ liệu giả lập và gán nhãn tham chiếu")
add_para(
"Dữ liệu giả lập được tạo theo phân bố hợp lý của các dấu hiệu sinh tồn theo từng mức độ (Xanh/Vàng/Đỏ). Nhãn tham chiếu được gán bằng "
"quy tắc dựa trên red flags và mức độ bất thường. Mục tiêu 200–500 ca; tỷ lệ Đỏ khoảng 15–25% để phản ánh thực tế và đủ dữ liệu học."
)
add_h2("5.3. Baseline rule-based (bắt buộc để so sánh)")
add_para("Baseline: Nếu có red flags → Đỏ; nếu không, tính điểm bất thường để phân Vàng/Xanh.")
add_h2("5.4. Mô hình AI ước lượng Risk score")
add_para(
"Sử dụng Logistic Regression (dễ giải thích) hoặc Random Forest (mạnh hơn và thuận lợi tính uncertainty). Đầu ra là xác suất nguy kịch p."
)
add_equation_block(
"Logistic Regression:\n"
"  p = P(critical) = σ(z) = 1 / (1 + e^{-z})\n"
"  z = β0 + β1·Age + β2·HR + β3·SBP + β4·SpO2 + β5·RR + β6·AVPU + ...\n"
"  Risk% = 100·p"
)
add_para("Gợi ý ngưỡng phân luồng (có thể tối ưu bằng dữ liệu):")
doc.add_paragraph("• Risk% ≥ 70% → gợi ý Đỏ")
doc.add_paragraph("• 30% ≤ Risk% < 70% → gợi ý Vàng")
doc.add_paragraph("• Risk% < 30% → gợi ý Xanh")
add_h2("5.5. Tính Uncertainty (độ không chắc chắn)")
add_para(
"Khuyến nghị: Random Forest. Mỗi cây t cho xác suất p_t. Uncertainty được tính bằng độ phân tán các p_t."
)
add_equation_block(
"Random Forest Uncertainty:\n"
"  p̄ = (1/T) · Σ_{t=1..T} p_t\n"
"  u = sqrt( (1/(T-1)) · Σ_{t=1..T} (p_t - p̄)^2 )\n"
"  Risk% = 100·p̄"
)
add_para("Phân mức tin cậy (ví dụ): u < 0.10: cao; 0.10–0.20: trung bình; ≥0.20: thấp (cần bác sĩ đánh giá kỹ).")
add_h2("5.6. Cơ chế Human-in-the-loop dựa trên Risk + Uncertainty")
add_para(
"Quy tắc khuyến nghị hành động: (i) Risk cao & u thấp → cảnh báo mạnh; (ii) Risk cao & u cao → cảnh báo + yêu cầu bác sĩ đánh giá kỹ; "
"(iii) Risk thấp & u thấp → ưu tiên thấp. Bác sĩ luôn có quyền bỏ qua khuyến nghị."
)
doc.add_paragraph()

add_h1("6. KIẾN TRÚC HỆ THỐNG")
add_equation_block(
"Pipeline:\n"
"  Input vitals/symptoms\n"
"     → Red-flag rules\n"
"     → Risk model (p)\n"
"     → Uncertainty (u)\n"
"     → Recommendation + explanation\n"
"     → Doctor final decision"
)
add_para("Hệ thống trình diễn (prototype) nhập nhanh dữ liệu, hiển thị Risk%, Uncertainty và lý do (feature importance/các dấu hiệu bất thường).")
doc.add_paragraph()

add_h1("7. THỰC NGHIỆM VÀ ĐÁNH GIÁ")
add_h2("7.1. Chia dữ liệu và huấn luyện")
add_para("Chia train/test (ví dụ 80/20) hoặc k-fold cross validation. Áp dụng class-weight để xử lý mất cân bằng lớp Đỏ.")
add_h2("7.2. Chỉ số đánh giá an toàn")
add_para("Đánh giá tập trung vào an toàn (nhóm Đỏ): Recall/Sensitivity, tỷ lệ bỏ sót nguy kịch, F1 theo lớp và confusion matrix.")
add_equation_block(
"Confusion Matrix (3 lớp):\n"
"  dự đoán vs thực tế (Đỏ/Vàng/Xanh)\n\n"
"Recall (Đỏ):\n"
"  Recall_red = TP_red / (TP_red + FN_red)\n\n"
"Miss rate (bỏ sót Đỏ):\n"
"  Miss_red = FN_red / (TP_red + FN_red) = 1 - Recall_red"
)
add_h2("7.3. Đánh giá tốc độ")
add_para("Đo thời gian suy luận/ca (ms) và tổng thời gian thao tác. Lưu ý: nhập liệu là phần bắt buộc dù có hay không có AI; AI giúp giảm thời gian do dự và giảm sửa sai.")
doc.add_paragraph()

add_h1("8. MINH HỌA TÌNH HUỐNG (CASE STUDY)")
add_h2("8.1. Case A: Risk cao, Uncertainty thấp → quyết nhanh")
add_para("Ví dụ: SpO₂ 88%, SBP 92, tri giác lơ mơ → Risk 76%, u thấp → cảnh báo mạnh, ưu tiên xử trí.")
add_h2("8.2. Case B: Risk cao, Uncertainty cao → bác sĩ vào cuộc")
add_para("Ví dụ: SpO₂ 93%, SBP 105, khó thở mơ hồ → Risk 74% nhưng u cao → hệ thống yêu cầu bác sĩ đánh giá kỹ.")
doc.add_paragraph()

add_h1("9. THẢO LUẬN")
add_para(
"Risk score giúp chuyển từ phân loại cứng sang định lượng nguy cơ, phù hợp quyết định y khoa. Uncertainty giúp mô hình an toàn hơn, tránh tự tin mù quáng "
"và kích hoạt Human-in-the-loop đúng lúc. Hạn chế: dữ liệu giả lập có thể thiếu đa dạng; cần thiết kế ca khó và kiểm soát thiên lệch."
)
doc.add_paragraph()

add_h1("10. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN")
add_para(
"Đề tài chứng minh hướng tiếp cận 'AI an toàn' cho phân luồng cấp cứu: định lượng rủi ro + độ không chắc chắn để hỗ trợ bác sĩ trong 30–60 giây đầu. "
"Trong tương lai: triển khai shadow mode, thu thập dữ liệu thực (ẩn danh), tinh chỉnh ngưỡng theo bối cảnh, mở rộng đặc trưng và đánh giá lâm sàng."
)

doc.add_paragraph()
add_h1("TÀI LIỆU THAM KHẢO (GỢI Ý)")
doc.add_paragraph("1. Tài liệu triage cơ bản (Bộ Y tế/WHO/ESI/ATS) – tùy nguồn bạn dùng.")
doc.add_paragraph("2. Tài liệu về Logistic Regression/Random Forest và uncertainty estimation.")
doc.add_paragraph("3. Tài liệu về decision support & human-in-the-loop trong y tế.")

out_path = "/mnt/data/BaoCao_15trang_TriageAI_Risk_Uncertainty.docx"
doc.save(out_path)
out_path
